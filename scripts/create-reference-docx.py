#!/usr/bin/env python3
"""Apply print-preview styles to Pandoc's official reference.docx."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


STAGES = ("layout", "body", "headings", "code-tables", "final")


def style_by_id(document: Document, style_id: str):
    for style in document.styles:
        if style.style_id == style_id:
            return style
    raise KeyError(f"reference.docx does not contain style {style_id!r}")


def apply_layout(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(22)
        section.right_margin = Mm(17)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def apply_body(document: Document) -> None:
    for style_id in ("Normal", "BodyText", "FirstParagraph"):
        style = style_by_id(document, style_id)
        style.font.name = "Liberation Serif"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    compact = style_by_id(document, "Compact")
    compact.font.name = "Liberation Serif"
    compact.font.size = Pt(10.5)
    compact.paragraph_format.space_after = Pt(2)
    compact.paragraph_format.line_spacing = 1.05


def apply_headings(document: Document) -> None:
    settings = {
        "Heading1": (16, 0, 12, True),
        "Heading2": (13, 14, 6, False),
        "Heading3": (11.5, 11, 5, False),
    }
    for style_id, (size, before, after, page_break) in settings.items():
        style = style_by_id(document, style_id)
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = page_break


def apply_code_and_tables(document: Document) -> None:
    try:
        source_code = style_by_id(document, "SourceCode")
    except KeyError:
        source_code = document.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    source_code.font.name = "Liberation Mono"
    source_code.font.size = Pt(8.5)
    source_code.paragraph_format.space_before = Pt(3)
    source_code.paragraph_format.space_after = Pt(5)
    source_code.paragraph_format.line_spacing = 1.0

    verbatim = style_by_id(document, "VerbatimChar")
    verbatim.font.name = "Liberation Mono"
    verbatim.font.size = Pt(8.5)

    table = style_by_id(document, "Table")
    table.font.name = "Liberation Sans"
    table.font.size = Pt(9)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def apply_page_numbers(document: Document) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        add_page_number(paragraph)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("base", type=Path, help="official Pandoc reference.docx")
    parser.add_argument("output", type=Path, help="styled reference.docx")
    parser.add_argument("--stage", choices=STAGES, default="final")
    args = parser.parse_args()

    document = Document(args.base)
    stage_index = STAGES.index(args.stage)
    apply_layout(document)
    if stage_index >= STAGES.index("body"):
        apply_body(document)
    if stage_index >= STAGES.index("headings"):
        apply_headings(document)
    if stage_index >= STAGES.index("code-tables"):
        apply_code_and_tables(document)
    if stage_index >= STAGES.index("final"):
        apply_page_numbers(document)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
